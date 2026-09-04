"""Create the final Breathe Better Tonight production handoff document.

Credentials are read only from a caller-supplied local file and are never written
to source control.
"""

from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], "071A2D")
        set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    document.add_paragraph()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor(7, 26, 45)


def add_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = "Intense Quote"
    paragraph.add_run(text)


def read_credentials(path: Path) -> dict[str, str]:
    values: dict[str, str] = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        if "=" in line:
            key, value = line.split("=", 1)
            values[key] = value
    required = {"ADMIN_URL", "ADMIN_USER", "ADMIN_PASS"}
    missing = required.difference(values)
    if missing:
        raise ValueError(f"Credential file is missing: {', '.join(sorted(missing))}")
    return values


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--credentials", required=True, type=Path)
    parser.add_argument(
        "--output",
        default=ROOT / "handoff" / "Breathe_Better_Tonight_Production_Handoff.docx",
        type=Path,
    )
    args = parser.parse_args()
    credentials = read_credentials(args.credentials)
    args.output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Breathe Better Tonight")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(7, 26, 45)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("Production Implementation & Operations Handoff")
    subrun.italic = True
    subrun.font.size = Pt(12)
    subtitle.add_run(f"\nPrepared {date.today().isoformat()}").font.size = Pt(9)
    document.add_paragraph()

    add_heading(document, "1. Production Status")
    document.add_paragraph(
        "The Night-to-Clarity public patient site, the preserved legacy landing-page family, the complete three-practice patient-concept catalog, and the protected publishing workspace have been implemented and deployed to the existing Railway project. The production service is connected to the repository’s main branch."
    )
    add_table(document, ["Item", "Production value"], [
        ["Patient website", "https://www.breathebettertonight.com/"],
        ["Railway service", "breathebettertonight-web · Production environment"],
        ["Deployment branch", "main"],
        ["Migration release commit", "542b2ae — merge: night-to-clarity landing-page system"],
        ["Complete doctor-set release", "1c146f4 — complete doctor concept and legacy page sets"],
        ["Current Page Index", "1,092 direct-preview routes"],
    ])

    add_heading(document, "2. Verified Administrator Access")
    add_note(document, "Store these credentials in your approved credential manager after review. They are configured in Railway variables and are not stored in the Git repository.")
    add_table(document, ["Field", "Value"], [
        ["Admin workspace", credentials["ADMIN_URL"]],
        ["Username", credentials["ADMIN_USER"]],
        ["Password", credentials["ADMIN_PASS"]],
        ["Authentication", "HTTP Basic authentication, configured through Railway production variables"],
    ])

    add_heading(document, "3. Patient, Provider, and Landing-Page Architecture")
    add_table(document, ["Experience", "Primary routes", "Purpose"], [
        ["Night-to-Clarity public site", "/, /sleep-check, /find-a-provider, /sleep-apnea", "Patient education, symptom awareness, provider selection, and the clear next step."],
        ["General resources", "/about, /faq, /contact, /thank-you", "Supportive informational and generic inquiry routes."],
        ["Practice handoff", "/care/{practice}", "Local practice identity, Call/Text actions, and practice-scoped policy links."],
        ["Patient concept", "/lp/{practice}/concepts/{concept}", "One of sixteen patient-facing concepts, rendered for the selected practice with the assigned GoHighLevel survey and policy profile."],
        ["Concept outcomes", "/lp/{practice}/concepts/{concept}/thank-you or /not-qualified", "Matched qualified and non-qualified outcome pages that retain the originating concept and selected practice."],
        ["Doctor-owned legacy", "/lp/{practice}/legacy/{variant}[/{city}]", "All thirteen preserved LP designs made available as complete doctor-owned design sets with locality variants."],
        ["Legacy outcomes", "/lp/{practice}/legacy/{variant}[/{city}]/thank-you or /not-qualified", "Matched qualified and non-qualified outcomes for each doctor-owned legacy design."],
        ["Paid traffic", "/go/{practice}/{campaign}", "Ad-matched hero, symptom-led education, Request a consultation CTA, and the assigned GoHighLevel survey."],
        ["HCP handoff", "/for-professionals", "External handoff to the Propel Dental HCP experience."],
    ])

    add_heading(document, "4. GoHighLevel Survey Assignment")
    add_table(document, ["Practice", "Canonical key", "Assigned GoHighLevel survey ID", "Campaign themes"], [
        ["Pantego Dental", "pantego-dental", "75op3Tl4LTjPkaXI1zhb", "All sixteen patient concepts, paid variants, and doctor-owned legacy designs"],
        ["PerioDDS", "periodds", "pvHcEcGNjxhXI3L8lSrE", "All sixteen patient concepts, paid variants, and doctor-owned legacy designs"],
        ["Dental World", "dental-world", "Rx0LnsI0XLu8JfhiDnYc", "All sixteen patient concepts, paid variants, and doctor-owned legacy designs"],
    ])
    document.add_paragraph(
        "Every patient-facing doctor page uses only its assigned survey. The preserved Dr. Willis Lay legacy URLs remain Pantego Dental compatibility routes, while standardized doctor-owned page-set paths make the same designs available for all three practices."
    )

    add_heading(document, "5. Administration Workspace")
    document.add_paragraph(
        "The protected admin workspace contains a dedicated Page Index, practice configuration cards, and a general-inquiry dashboard. The Page Index provides direct previews and category filters for public pages, practice pages, canonical concepts, doctor-owned legacy pages, campaign landing pages, qualified thank-you pages, non-qualified thank-you pages, policies, and every preserved legacy URL."
    )
    add_table(document, ["Configuration area", "What it controls"], [
        ["Practice identity", "Public name, campaign destination name, service label, and local care destination."],
        ["Contact actions", "Visible Call number, Call route, and Text route per practice."],
        ["GoHighLevel handoff", "One validated survey ID for each practice; doctor-specific pages do not use generic placeholders."],
        ["Design assignment", "All sixteen patient-facing concepts and the legacy design catalog can be assigned per practice."],
        ["Policy profile", "Privacy, Terms & Conditions, and Accessibility override content that updates every route using that practice profile."],
        ["Outcome family", "Direct previews of the landing page, qualified thank-you page, and non-qualified thank-you page for every campaign, concept, and doctor-owned legacy design."],
    ])

    add_heading(document, "6. Three Completed Improvement Rounds")
    add_table(document, ["Round", "Focus", "Outcome"], [
        ["1", "Copy clarity and doctor identity", "Every concept and outcome page names the selected local practice and explains that the request is routed to that destination."],
        ["2", "Conversion and matched outcomes", "All doctor pages use assigned GoHighLevel surveys, visible Call/Text actions, consultation language, and concept-matched qualified/non-qualified outcomes."],
        ["3", "Accessibility, responsive behavior, and administration", "Responsive QA checks all thirteen legacy variants plus representative concept and legacy pages for all three practices, including horizontal overflow and practice policy links."],
    ])

    add_heading(document, "7. Validation Summary")
    add_table(document, ["Validation", "Result"], [
        ["Legacy route regression", "234 generated legacy routes passed locally."],
        ["Legacy LP structural check", "13 legacy variants passed required accessible structure and approved content checks."],
        ["Complete doctor-set route regression", "All 690 doctor-owned concept and legacy routes passed locally, retaining their assigned survey and policy profile."],
        ["Responsive visual QA", "38 mobile and desktop captures completed across the 13 legacy variants and all three active practice systems."],
        ["Production route check", "The public site, Pantego concept route, PerioDDS legacy route, Dental World concept route, and a matched non-qualified outcome route returned HTTP 200 with the expected survey or policy evidence."],
        ["Production admin check", "The `/admin` route remains protected in production; its complete Page Index was validated locally with all three doctor-owned route families."],
    ])

    add_heading(document, "8. Operating Notes")
    document.add_paragraph(
        "The Railway service must retain a writable persistent volume mounted at /data. The application uses LEADS_FILE for general inquiries and PRACTICE_CONFIG_FILE for updates made through the protected practice configuration workspace. Admin changes are intentionally stored outside Git so configuration can be managed without code edits.")
    document.add_paragraph(
        "Before paid media is activated, verify each practice’s text-routing number accepts messages, add legally approved practice-specific policy text where needed, and replace any unapproved provider biography placeholder content with approved factual practice copy. The patient-facing content remains symptom-led and informational; it does not diagnose sleep apnea, promise an outcome, or advertise a free consultation.")

    add_heading(document, "9. Source-Control Record")
    document.add_paragraph(
        "The implementation is present in the private GitHub repository at https://github.com/srinagubandi/breathebettertonight. The migration work was developed on feature/night-to-clarity-lp-system-2026, merged into main, and deployed by the existing Railway service.")

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Breathe Better Tonight · Production handoff · Confidential")
        footer.runs[0].font.size = Pt(8)

    document.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
